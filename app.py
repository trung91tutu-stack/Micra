"""Giao diện web MICRA (Streamlit) — bản 3.

    streamlit run app.py

Chín màn hình:
    Giới thiệu · Tổng quan · Thẩm định hồ sơ · Nhập hồ sơ mới · Nạp dữ liệu mới
    So sánh hồ sơ · Giám sát sau vay · Phỏng vấn qua Zalo · Kiểm chứng mô hình

Nguyên tắc không đổi: mô hình ngôn ngữ chỉ diễn đạt, mọi con số đều do tầng
thống kê tất định tính ra.
"""
from __future__ import annotations
import hashlib
import io
import shutil
import sys
from datetime import datetime
from pathlib import Path

import altair as alt
import numpy as np
import pandas as pd
import streamlit as st

GOC = Path(__file__).resolve().parent
sys.path.insert(0, str(GOC / "src"))
from schema import HoSo
from orchestrator import tham_dinh
from features import TEN_DAC_TRUNG, DIEN_CHUAN_NGANH, dien_giai, COT
from llm import dang_dung, goi_llm
from agents.a2_phong_van import HE_THONG, CAU_HOI
from agents import a6_giam_sat
from mo_phong_sau_vay import ba_thang_tiep
import scoring
from scoring import NGUONG

_V = tuple(int(x) for x in st.__version__.split(".")[:2] if x.isdigit())
ROI = {"width": "stretch"} if _V >= (1, 49) else {"use_container_width": True}

DATA = GOC / "data" / "du_lieu_mo_phong_60_ho_kinh_doanh.csv"
ASSETS = GOC / "assets"
MODELS = GOC / "models"
LOGO = ASSETS / "MICRA_lockup_doc.png"
ICON = ASSETS / "MICRA_favicon_32.png"

st.set_page_config(page_title="MICRA — Thẩm định tín dụng hộ kinh doanh",
                   page_icon=str(ICON) if ICON.exists() else "🏦",
                   layout="wide", initial_sidebar_state="expanded")

TEAL, DARK, SEA, AMBER = "#035E63", "#0B2B2E", "#00A896", "#E8963C"
INK, GREY, LIGHT, BAND = "#17302F", "#5E7270", "#F4F8F7", "#E4F0EE"
MAU = {"Thấp": "#1F7A3D", "Trung bình": "#B8860B", "Cao": "#D2691E", "Rất cao": "#C00000"}

st.markdown(f"""
<style>
  .block-container {{ padding-top: 2.1rem; padding-bottom: 3rem; max-width: 1500px; }}
  h1, h2, h3 {{ color: {INK}; letter-spacing: -.01em; }}
  h1 {{ font-size: 1.9rem !important; font-weight: 700 !important; }}

  .micra-sub {{ font-size: .74rem; color: {GREY}; letter-spacing: .04em;
                text-align: center; margin: -.3rem 0 .7rem 0; }}

  .kpi {{ background: {LIGHT}; border: 1px solid #DCE8E6; border-radius: 10px;
          padding: .8rem .95rem; height: 100%; }}
  .kpi .lab {{ font-size: .74rem; color: {GREY}; text-transform: uppercase;
               letter-spacing: .06em; margin-bottom: .28rem; }}
  .kpi .val {{ font-size: 1.6rem; font-weight: 700; color: {INK}; line-height: 1.1; }}
  .kpi .sub {{ font-size: .74rem; color: {GREY}; margin-top: .22rem; }}
  .kpi.hi {{ background: #FDF3E8; border-color: #F0D3AE; }}
  .kpi.hi .val {{ color: #A0641F; }}

  .band {{ border-radius: 10px; padding: .85rem 1.1rem; margin: .2rem 0 .9rem 0;
           color: #fff; background: {DARK}; font-size: .92rem; }}
  .note {{ background: {BAND}; border-left: 5px solid {TEAL}; border-radius: 8px;
           padding: .7rem .95rem; font-size: .86rem; color: {INK}; }}
  .warn {{ background: #FDF3E8; border-left: 5px solid {AMBER}; border-radius: 8px;
           padding: .7rem .95rem; font-size: .86rem; color: {INK}; }}

  .chip {{ display: inline-block; padding: .18rem .6rem; border-radius: 999px;
           font-size: .74rem; font-weight: 600; color: #fff; margin-right: .3rem; }}

  .stepbar {{ display: flex; gap: 6px; margin: .1rem 0 .9rem 0; }}
  .step {{ flex: 1; border-radius: 6px; padding: .42rem .3rem; text-align: center;
           font-size: .72rem; font-weight: 600; border: 1px solid #DCE8E6;
           background: {LIGHT}; color: {GREY}; }}
  .step.on {{ background: {TEAL}; color: #fff; border-color: {TEAL}; }}
  .step.llm {{ background: #FDF3E8; color: #A0641F; border-color: #F0D3AE; }}

  .tacTu {{ border: 1px solid #DCE8E6; border-radius: 10px; padding: .75rem .9rem;
            background: #fff; height: 100%; }}
  .tacTu .num {{ display: inline-block; width: 1.5rem; height: 1.5rem; border-radius: 50%;
                 background: {TEAL}; color: #fff; text-align: center; line-height: 1.5rem;
                 font-size: .8rem; font-weight: 700; margin-right: .4rem; }}
  .tacTu.llm .num {{ background: {AMBER}; }}
  .tacTu .ten {{ font-weight: 700; color: {INK}; font-size: .95rem; }}
  .tacTu .mo {{ font-size: .78rem; color: {GREY}; margin-top: .4rem; line-height: 1.45; }}

  div[data-testid="stMetricValue"] {{ font-size: 1.45rem; }}
  section[data-testid="stSidebar"] {{ background: #FBFDFD; }}
  #MainMenu, footer {{ visibility: hidden; }}
</style>
""", unsafe_allow_html=True)


# ─────────────────────────────────────────────── tiện ích hiển thị
def kpi(col, nhan, gia_tri, phu="", hi=False):
    col.markdown(f"<div class='kpi{' hi' if hi else ''}'><div class='lab'>{nhan}</div>"
                 f"<div class='val'>{gia_tri}</div><div class='sub'>{phu}&nbsp;</div></div>",
                 unsafe_allow_html=True)


def chip(text, mau):
    return f"<span class='chip' style='background:{mau}'>{text}</span>"


def thanh_buoc(den=6, llm_o=4):
    ten = ["1 Thu thập", "2 Phỏng vấn", "3 Phân tích", "4 Tờ trình", "5 Kiểm soát", "6 Giám sát"]
    o = "".join(f"<div class='step {'llm' if (i+1) == llm_o and (i+1) <= den else ('on' if i < den else '')}'>{t}</div>"
                for i, t in enumerate(ten))
    st.markdown(f"<div class='stepbar'>{o}</div>", unsafe_allow_html=True)


def tien(x, don_vi=" tr"):
    return f"{x/1e6:,.0f}{don_vi}"


# ─────────────────────────────────────────────── biểu đồ Altair
TRUC = dict(labelColor=GREY, titleColor=GREY, labelFontSize=11, titleFontSize=11,
            gridColor="#EAF1F0", domainColor="#D6E4E2", tickColor="#D6E4E2")


def _khung(c, cao=280):
    return (c.properties(height=cao)
             .configure_view(strokeWidth=0)
             .configure(background="transparent")
             .configure_axis(**TRUC)
             .configure_legend(labelColor=GREY, titleColor=GREY, labelFontSize=11,
                               titleFontSize=11, orient="top", direction="horizontal"))


def bd_cot(df, truc_x, truc_y, mau=TEAL, cao=260, ngang=False, dinh_dang=".2f", nhan=True):
    b = alt.Chart(df).mark_bar(cornerRadius=3, color=mau)
    if ngang:
        b = b.encode(y=alt.Y(f"{truc_x}:N", sort="-x", title=None),
                     x=alt.X(f"{truc_y}:Q", title=None))
        t = alt.Chart(df).mark_text(align="left", dx=4, color=GREY, fontSize=11).encode(
            y=alt.Y(f"{truc_x}:N", sort="-x"), x=f"{truc_y}:Q",
            text=alt.Text(f"{truc_y}:Q", format=dinh_dang))
    else:
        b = b.encode(x=alt.X(f"{truc_x}:N", sort=None, title=None),
                     y=alt.Y(f"{truc_y}:Q", title=None))
        t = alt.Chart(df).mark_text(dy=-8, color=GREY, fontSize=11).encode(
            x=alt.X(f"{truc_x}:N", sort=None), y=f"{truc_y}:Q",
            text=alt.Text(f"{truc_y}:Q", format=dinh_dang))
    return _khung((b + t) if nhan else b, cao)


def bd_cot_am_duong(df, truc_x, truc_y, cao=340):
    b = alt.Chart(df).mark_bar(cornerRadius=3).encode(
        y=alt.Y(f"{truc_x}:N", sort=alt.EncodingSortField(field=truc_y, order="descending"), title=None),
        x=alt.X(f"{truc_y}:Q", title="Đóng góp vào điểm rủi ro"),
        color=alt.condition(alt.datum[truc_y] > 0, alt.value("#C0703A"), alt.value(TEAL)),
        tooltip=[truc_x, alt.Tooltip(f"{truc_y}:Q", format="+.3f")])
    return _khung(b, cao)


def bd_duong(df, truc_x, cot, cao=290, mau=None):
    d = df.melt(truc_x, var_name="Chỉ tiêu", value_name="Giá trị")
    thang = alt.Scale(range=mau) if mau else alt.Scale(scheme="tableau10")
    c = alt.Chart(d).mark_line(point=alt.OverlayMarkDef(size=45), strokeWidth=2.4).encode(
        x=alt.X(f"{truc_x}:N", sort=None, title=None),
        y=alt.Y("Giá trị:Q", title=None),
        color=alt.Color("Chỉ tiêu:N", scale=thang, title=None),
        tooltip=[truc_x, "Chỉ tiêu", alt.Tooltip("Giá trị:Q", format=",.1f")])
    return _khung(c, cao)


def bd_roc(roc, cao=300):
    duong = alt.Chart(roc).mark_line(color=TEAL, strokeWidth=2.6).encode(
        x=alt.X("Tỷ lệ báo động giả:Q", scale=alt.Scale(domain=[0, 1])),
        y=alt.Y("Tỷ lệ bắt đúng:Q", scale=alt.Scale(domain=[0, 1])))
    cheo = alt.Chart(pd.DataFrame({"x": [0, 1], "y": [0, 1]})).mark_line(
        color=GREY, strokeDash=[5, 4], strokeWidth=1.2).encode(x="x:Q", y="y:Q")
    return _khung(cheo + duong, cao)


def bd_hieu_chinh(hc, cao=300):
    cheo = alt.Chart(pd.DataFrame({"x": [0, 1], "y": [0, 1]})).mark_line(
        color=GREY, strokeDash=[5, 4], strokeWidth=1.2).encode(
        x=alt.X("x:Q", title="Xác suất dự báo", scale=alt.Scale(domain=[0, 1])),
        y=alt.Y("y:Q", title="Tỷ lệ vỡ nợ thực tế", scale=alt.Scale(domain=[0, 1])))
    d = alt.Chart(hc).mark_line(point=alt.OverlayMarkDef(size=70, color=AMBER),
                                color=AMBER, strokeWidth=2.4).encode(
        x="Xác suất dự báo:Q", y="Tỷ lệ vỡ nợ thực tế:Q")
    return _khung(cheo + d, cao)


# ─────────────────────────────────────────────── dữ liệu
COT_THANG = [f"{p}_T{i}" for p in ("doanh_thu_hoa_don", "dong_tien_ngan_hang", "tien_dien")
             for i in range(1, 13)]
COT_BAT_BUOC = (["ma_ho", "ten_ho", "nganh", "tinh_thanh", "so_nam_hoat_dong",
                 "so_giao_dich_trung_binh_ngay", "co_lich_su_tin_dung_CIC",
                 "so_tien_de_nghi_vay"] + COT_THANG)


@st.cache_data
def nap_mac_dinh():
    return pd.read_csv(DATA)


def khoa_bang(df) -> str:
    return hashlib.md5(pd.util.hash_pandas_object(df, index=False).values).hexdigest()[:16]


def bang_hien_hanh() -> pd.DataFrame:
    return st.session_state.get("bang_moi", nap_mac_dinh())


def ten_nguon() -> str:
    return st.session_state.get("ten_nguon", "Bộ mô phỏng 60 hộ (mặc định)")


def kiem_tra_bang(df):
    """Trả về (danh sách lỗi, danh sách cảnh báo)."""
    loi, canh_bao = [], []
    thieu = [c for c in COT_BAT_BUOC if c not in df.columns]
    if thieu:
        loi.append(f"Thiếu {len(thieu)} cột bắt buộc: " + ", ".join(thieu[:8])
                   + (" …" if len(thieu) > 8 else ""))
        return loi, canh_bao
    if len(df) < 5:
        loi.append(f"Chỉ có {len(df)} dòng. Cần ít nhất 5 hồ sơ.")
    if df["ma_ho"].duplicated().any():
        loi.append("Cột ma_ho có giá trị trùng lặp.")
    so = ["so_nam_hoat_dong", "so_giao_dich_trung_binh_ngay", "so_tien_de_nghi_vay"] + COT_THANG
    for c in so:
        if not pd.api.types.is_numeric_dtype(df[c]):
            loi.append(f"Cột {c} phải là số, đang là {df[c].dtype}.")
            break
    if not loi:
        if df[so].isna().any().any():
            loi.append("Có ô trống trong các cột số. Điền đủ trước khi nạp.")
        if (df[so] < 0).any().any():
            canh_bao.append("Có giá trị âm trong các cột số — kiểm tra lại đơn vị.")
    la = set(df["nganh"].dropna().unique()) - set(DIEN_CHUAN_NGANH)
    if la:
        canh_bao.append("Ngành chưa có chuẩn tiền điện, hệ thống dùng mức chung 2,5%: "
                        + ", ".join(sorted(la)[:6]))
    hop_le_cic = set(df["co_lich_su_tin_dung_CIC"].astype(str).str.strip().unique()) - {"Có", "Không"}
    if hop_le_cic:
        canh_bao.append("Cột co_lich_su_tin_dung_CIC nên chỉ chứa 'Có' hoặc 'Không'. "
                        "Giá trị khác được hiểu là 'Không'.")
    if "nhan_vo_no" not in df.columns:
        canh_bao.append("Không có cột nhan_vo_no nên chỉ chấm điểm được, "
                        "không huấn luyện lại được mô hình.")
    return loi, canh_bao


@st.cache_resource(show_spinner=False)
def _chay_mot(khoa: str, ma: str, dung_llm: bool):
    df = bang_hien_hanh()
    return tham_dinh(HoSo.tu_dong_csv(df[df.ma_ho == ma].iloc[0].to_dict()), dung_llm=dung_llm)


@st.cache_resource(show_spinner="Đang chấm điểm toàn bộ hồ sơ…")
def _chay_tat_ca(khoa: str):
    df = bang_hien_hanh()
    return {r["ma_ho"]: tham_dinh(HoSo.tu_dong_csv(r), dung_llm=False)
            for r in df.to_dict("records")}


def chay_agent(ma, dung_llm):
    return _chay_mot(st.session_state["khoa"], ma, dung_llm)


def chay_tat_ca():
    return _chay_tat_ca(st.session_state["khoa"])


@st.cache_data(show_spinner="Đang chạy kiểm định chéo…")
def chi_so_mo_hinh(khoa: str):
    from sklearn.linear_model import LogisticRegression
    from sklearn.preprocessing import StandardScaler
    from sklearn.pipeline import make_pipeline
    from sklearn.model_selection import StratifiedKFold, cross_val_predict
    from sklearn.calibration import calibration_curve
    from sklearn.metrics import (roc_auc_score, average_precision_score,
                                 brier_score_loss, roc_curve)
    from train import nap_tu_bang

    df = bang_hien_hanh()
    if "nhan_vo_no" not in df.columns:
        return None
    X, y, _ = nap_tu_bang(df)
    ong = lambda: make_pipeline(StandardScaler(), LogisticRegression(max_iter=5000, C=0.6))
    aucs, prs, brs = [], [], []
    for lap in range(10):
        cv = StratifiedKFold(5, shuffle=True, random_state=lap)
        p = cross_val_predict(ong(), X, y, cv=cv, method="predict_proba")[:, 1]
        aucs.append(roc_auc_score(y, p)); prs.append(average_precision_score(y, p))
        brs.append(brier_score_loss(y, p))
    p0 = cross_val_predict(ong(), X, y, cv=StratifiedKFold(5, shuffle=True, random_state=0),
                           method="predict_proba")[:, 1]
    rng = np.random.default_rng(0); boot = []
    for _ in range(2000):
        b = rng.choice(len(y), len(y), replace=True)
        if len(np.unique(y[b])) > 1:
            boot.append(roc_auc_score(y[b], p0[b]))
    fpr, tpr, _ = roc_curve(y, p0)
    nbin = min(5, max(2, int(y.sum())))
    tx, dx = calibration_curve(y, p0, n_bins=nbin, strategy="quantile")
    return {"auc": (float(np.mean(aucs)), float(np.std(aucs))),
            "pr": (float(np.mean(prs)), float(np.std(prs))),
            "brier": (float(np.mean(brs)), float(np.std(brs))),
            "ktc": (float(np.percentile(boot, 2.5)), float(np.percentile(boot, 97.5))),
            "roc": pd.DataFrame({"Tỷ lệ báo động giả": fpr, "Tỷ lệ bắt đúng": tpr}),
            "hieu_chinh": pd.DataFrame({"Xác suất dự báo": dx, "Tỷ lệ vỡ nợ thực tế": tx}),
            "n": int(len(y)), "n_vo_no": int(y.sum()), "p_cv": p0, "y": y}


def to_docx(tt):
    try:
        from docx import Document
        from docx.shared import Pt
    except Exception:
        return None
    d = Document()
    d.add_heading(f"TỜ TRÌNH THẨM ĐỊNH — {tt.ho_so.ten_ho} ({tt.ho_so.ma_ho})", level=1)
    for dong in tt.to_trinh.split("\n"):
        p = d.add_paragraph(dong)
        if dong.isupper() and dong.strip() and p.runs:
            p.runs[0].bold = True
        for r in p.runs:
            r.font.size = Pt(11)
    d.add_page_break()
    d.add_heading("DANH SÁCH BẰNG CHỨNG", level=2)
    b = d.add_table(rows=1, cols=3)
    b.style = "Light Grid Accent 1"
    for i, h in enumerate(["Mã", "Nội dung", "Nguồn"]):
        b.rows[0].cells[i].text = h
    for bc in tt.bang_chung:
        c = b.add_row().cells
        c[0].text, c[1].text, c[2].text = bc.ma, bc.noi_dung, bc.nguon
    buf = io.BytesIO(); d.save(buf)
    return buf.getvalue()


# ─────────────────────────────────────────────── thanh bên
df = bang_hien_hanh()
st.session_state["khoa"] = khoa_bang(df)

with st.sidebar:
    if LOGO.exists():
        st.image(str(LOGO), **ROI)
    else:
        st.markdown("<div style='font-size:1.6rem;font-weight:800;color:%s;"
                    "letter-spacing:.16em;text-align:center'>MICRA</div>" % TEAL,
                    unsafe_allow_html=True)

    man = st.radio("Màn hình",
                   ["Giới thiệu", "Tổng quan", "Thẩm định hồ sơ", "Nhập hồ sơ mới",
                    "Nạp dữ liệu mới", "So sánh hồ sơ", "Giám sát sau vay",
                    "Phỏng vấn qua Zalo", "Kiểm chứng mô hình"],
                   label_visibility="collapsed")

    st.divider()
    nha = dang_dung()
    st.caption(f"Bộ dữ liệu: **{ten_nguon()}** · {len(df)} hồ sơ")
    st.caption(f"Mô hình ngôn ngữ: **{nha}**")
    if nha == "mock":
        st.info("Chế độ mô phỏng. Thêm khóa API vào .env để bật mô hình ngôn ngữ thật. "
                "Mọi con số vẫn do tầng thống kê tính, không đổi.")
    dung_llm_ct = st.toggle("Dùng mô hình ngôn ngữ khi soạn tờ trình", value=(nha != "mock"),
                            help="Tắt đi thì tác tử 4 soạn tờ trình bằng quy tắc tất định — "
                                 "nhanh hơn, không tốn hạn mức. Con số không đổi.")

# ══════════════════════════════════════════════ 0. GIỚI THIỆU
if man == "Giới thiệu":
    a, b_ = st.columns([1, 2.4])
    with a:
        if LOGO.exists():
            st.image(str(LOGO), **ROI)
    with b_:
        st.markdown("# Thẩm định tín dụng hộ kinh doanh bằng hệ thống sáu tác tử AI")
        st.markdown(
            "<div class='note'>Khoảng 5 triệu hộ kinh doanh cá thể bị loại khỏi tín dụng "
            "chính thức. Không phải vì họ quá rủi ro, mà vì <b>chi phí thẩm định gần như "
            "không đổi theo quy mô khoản vay</b> — thẩm định khoản 80 triệu tốn gần bằng "
            "khoản 5 tỉ. Mọi khoản dưới ngưỡng C/m bị từ chối vì không đủ bù chi phí vận "
            "hành. MICRA hạ C để hạ ngưỡng đó.</div>", unsafe_allow_html=True)
        st.write("")
        c = st.columns(4)
        kpi(c[0], "Năng lực phân biệt", "0,858", "ROC-AUC, kiểm định chéo", hi=True)
        kpi(c[1], "Tỷ lệ truy vết số", "100%", "mọi con số về được bằng chứng")
        kpi(c[2], "Kiểm thử tự động", "16/16", "mục đạt")
        kpi(c[3], "Tốc độ", "≈ 35 ms", "mỗi hồ sơ, chế độ tất định")

    st.divider()
    st.markdown("### Hành trình một hồ sơ qua sáu tác tử")
    A6 = [("1", "Thu thập", "Hóa đơn điện tử, sao kê ngân hàng, tiền điện, CIC — sau khi chủ hộ đồng ý.", False),
          ("2", "Phỏng vấn", "Hỏi qua Zalo bằng tiếng Việt đời thường, thu thông tin mềm thay khảo sát thực địa.", True),
          ("3", "Phân tích", "11 đặc trưng tất định → mô hình thống kê → xác suất vỡ nợ. Không dùng AI ngôn ngữ.", False),
          ("4", "Tờ trình", "AI ngôn ngữ diễn đạt dữ kiện thành văn bản cho cán bộ đọc trong 5 phút.", True),
          ("5", "Kiểm soát", "Hai lớp: truy vết số tất định và rà soát ngữ nghĩa. Chặn mọi con số không nguồn.", False),
          ("6", "Giám sát", "Chạy hằng tuần sau giải ngân, 5 luật cảnh báo sớm trước khi thành nợ xấu.", False)]
    cot = st.columns(6)
    for c, (n, ten, mo, llm) in zip(cot, A6):
        c.markdown(f"<div class='tacTu{' llm' if llm else ''}'><span class='num'>{n}</span>"
                   f"<span class='ten'>{ten}</span><div class='mo'>{mo}</div></div>",
                   unsafe_allow_html=True)
    st.write("")
    st.markdown(f"<div class='band' style='background:{AMBER};color:{DARK};font-weight:600'>"
                "QUY TẮC VÀNG — Tầng thống kê quyết định MỌI CON SỐ. Mô hình ngôn ngữ chỉ xử lý "
                "ngôn ngữ. Nhờ vậy hệ thống giải trình được, ổn định và kiểm soát được thiên lệch."
                "</div>", unsafe_allow_html=True)

    st.divider()
    g1, g2 = st.columns(2)
    g1.markdown("#### Bắt đầu từ đâu")
    g1.markdown("- **Tổng quan** — nhìn toàn bộ danh mục trong một màn hình\n"
                "- **Thẩm định hồ sơ** — đi sâu vào một hồ sơ cụ thể\n"
                "- **Nhập hồ sơ mới** — tự nhập số liệu để hệ thống chấm điểm\n"
                "- **Nạp dữ liệu mới** — tải bộ hồ sơ của bạn lên và huấn luyện lại mô hình")
    g2.markdown("#### Điều cần biết")
    g2.markdown(f"<div class='warn'>Bộ dữ liệu đang dùng là <b>{ten_nguon()}</b>. "
                "Dữ liệu mô phỏng không phải thông tin thật của bất kỳ cá nhân hay tổ chức nào. "
                "MICRA không cho vay, không giữ tiền, không chịu rủi ro tín dụng — quyết định "
                "phê duyệt luôn thuộc về cán bộ tín dụng.</div>", unsafe_allow_html=True)

# ══════════════════════════════════════════════ 1. TỔNG QUAN
elif man == "Tổng quan":
    st.title("Tổng quan danh mục")
    st.caption(f"Nguồn: {ten_nguon()} · chấm điểm bằng tầng thống kê tất định, không gọi mô hình ngôn ngữ.")

    tat_ca = chay_tat_ca()
    bang = pd.DataFrame([{
        "Mã": m, "Hộ kinh doanh": t.ho_so.ten_ho, "Ngành": t.ho_so.nganh,
        "Tỉnh/TP": t.ho_so.tinh_thanh, "Số năm": round(t.ho_so.so_nam_hoat_dong, 1),
        "DT TB (tr)": round(t.dac_trung["dt_trung_binh"] / 1e6),
        "Đề nghị (tr)": round(t.ho_so.so_tien_de_nghi_vay / 1e6),
        "Xác suất vỡ nợ": t.xac_suat_vo_no, "Hạng": t.hang_rui_ro,
        "Đề xuất duyệt (tr)": round(t.khuyen_nghi["de_xuat_duyet"] / 1e6),
        "CIC": "Có" if t.ho_so.co_cic else "Không"} for m, t in tat_ca.items()])

    c = st.columns(5)
    kpi(c[0], "Tổng hồ sơ", f"{len(bang)}", f"{bang['Ngành'].nunique()} ngành nghề")
    kpi(c[1], "Dư nợ đề nghị", f"{bang['Đề nghị (tr)'].sum():,.0f} tr",
        f"trung bình {bang['Đề nghị (tr)'].mean():,.0f} tr/hồ sơ")
    kpi(c[2], "Đề xuất duyệt", f"{bang['Đề xuất duyệt (tr)'].sum():,.0f} tr",
        f"{bang['Đề xuất duyệt (tr)'].sum()/max(bang['Đề nghị (tr)'].sum(),1):.0%} so với đề nghị")
    kpi(c[3], "Rủi ro cao trở lên", f"{int(bang['Hạng'].isin(['Cao','Rất cao']).sum())}",
        f"{bang['Hạng'].isin(['Cao','Rất cao']).mean():.0%} danh mục", hi=True)
    kpi(c[4], "Xác suất vỡ nợ TB", f"{bang['Xác suất vỡ nợ'].mean():.1%}",
        f"trung vị {bang['Xác suất vỡ nợ'].median():.1%}")
    st.write("")

    g1, g2 = st.columns([2, 3])
    with g1:
        st.markdown("**Phân bố theo hạng rủi ro**")
        tu = ["Thấp", "Trung bình", "Cao", "Rất cao"]
        d = (bang["Hạng"].value_counts().reindex(tu).fillna(0).astype(int)
             .rename_axis("Hạng").reset_index(name="Số hồ sơ"))
        c1 = alt.Chart(d).mark_bar(cornerRadius=3).encode(
            x=alt.X("Hạng:N", sort=tu, title=None),
            y=alt.Y("Số hồ sơ:Q", title=None),
            color=alt.Color("Hạng:N", sort=tu, legend=None,
                            scale=alt.Scale(domain=tu, range=[MAU[k] for k in tu])),
            tooltip=["Hạng", "Số hồ sơ"])
        t1 = alt.Chart(d).mark_text(dy=-9, color=GREY, fontSize=12).encode(
            x=alt.X("Hạng:N", sort=tu), y="Số hồ sơ:Q", text="Số hồ sơ:Q")
        st.altair_chart(_khung(c1 + t1, 250), **ROI)
    with g2:
        st.markdown("**Xác suất vỡ nợ trung bình theo ngành**")
        d2 = (bang.groupby("Ngành")["Xác suất vỡ nợ"].mean().mul(100)
              .round(1).reset_index(name="Tỷ lệ %"))
        st.altair_chart(bd_cot(d2, "Ngành", "Tỷ lệ %", SEA, 250, ngang=True, dinh_dang=".1f"), **ROI)

    st.divider()
    f1, f2, f3 = st.columns([2, 2, 3])
    ln = f1.multiselect("Lọc ngành", sorted(bang["Ngành"].unique()))
    lh = f2.multiselect("Lọc hạng rủi ro", ["Thấp", "Trung bình", "Cao", "Rất cao"])
    tim = f3.text_input("Tìm theo tên hoặc mã hồ sơ", placeholder="ví dụ: HKD023 hoặc Tân Tiến")
    b = bang.copy()
    if ln: b = b[b["Ngành"].isin(ln)]
    if lh: b = b[b["Hạng"].isin(lh)]
    if tim:
        k = tim.strip().lower()
        b = b[b["Mã"].str.lower().str.contains(k) | b["Hộ kinh doanh"].str.lower().str.contains(k)]
    st.caption(f"Hiển thị {len(b)} trên {len(bang)} hồ sơ")
    st.dataframe(b, hide_index=True, height=420, **ROI,
                 column_config={"Xác suất vỡ nợ": st.column_config.ProgressColumn(
                     "Xác suất vỡ nợ", format="%.1f%%", min_value=0.0, max_value=1.0)})
    st.download_button("Tải bảng kết quả (.csv)", b.to_csv(index=False).encode("utf-8-sig"),
                       "micra_ket_qua_tham_dinh.csv", "text/csv")

# ══════════════════════════════════════════════ 2. THẨM ĐỊNH
elif man == "Thẩm định hồ sơ":
    st.title("Thẩm định hồ sơ")
    ma = st.selectbox("Chọn hồ sơ", df.ma_ho.tolist(),
                      format_func=lambda m: f"{m} — {df[df.ma_ho==m].ten_ho.iloc[0]}")
    o = st.empty()
    with o.container():
        thanh_buoc(0); st.caption("Đang chạy sáu tác tử…")
    tt = chay_agent(ma, dung_llm_ct)
    o.empty()

    h, d = tt.ho_so, tt.dac_trung
    thanh_buoc(6, llm_o=4 if dung_llm_ct else 0)
    st.markdown(f"### {h.ten_ho}  " + chip(h.nganh, TEAL) + chip(h.tinh_thanh, GREY)
                + chip(f"Hạng {tt.hang_rui_ro}", MAU[tt.hang_rui_ro]), unsafe_allow_html=True)

    c = st.columns(5)
    kpi(c[0], "Số năm hoạt động", f"{h.so_nam_hoat_dong:.1f}", "năm")
    kpi(c[1], "Doanh thu TB", tien(d["dt_trung_binh"]), "triệu đồng/tháng")
    kpi(c[2], "Đề nghị vay", tien(h.so_tien_de_nghi_vay), f"{d['vay_tren_dt']:.2f}× doanh thu tháng")
    kpi(c[3], "Xác suất vỡ nợ", f"{tt.xac_suat_vo_no:.1%}", "trong 12 tháng tới", hi=True)
    kpi(c[4], "Đề xuất duyệt", tien(tt.khuyen_nghi["de_xuat_duyet"]),
        f"trần theo hạng {tien(tt.khuyen_nghi['tran_theo_hang'])}")
    st.write("")
    st.markdown(f"<div class='band' style='background:{MAU[tt.hang_rui_ro]}'>"
                f"<b>Khuyến nghị của hệ thống:</b> {tt.khuyen_nghi['quyet_dinh']}. "
                f"Ngưỡng phân hạng: {' · '.join(f'{n:.0%} → {t}' for n, t in NGUONG)} · còn lại Rất cao."
                "</div>", unsafe_allow_html=True)

    t1, t2, t3, t4 = st.tabs(["Tờ trình", "Dữ liệu 12 tháng", "Yếu tố ảnh hưởng", "Kiểm soát & nhật ký"])
    with t1:
        a, b_ = st.columns([3, 1])
        a.text_area("to_trinh", tt.to_trinh, height=480, label_visibility="collapsed")
        with b_:
            st.markdown("<div class='note'>Tờ trình do tác tử 4 soạn. Quyết định phê duyệt "
                        "cuối cùng thuộc về cán bộ tín dụng.</div>", unsafe_allow_html=True)
            st.write("")
            st.download_button("Tải tờ trình (.txt)", tt.to_trinh.encode("utf-8"),
                               f"to_trinh_{h.ma_ho}.txt", **ROI)
            bb = to_docx(tt)
            if bb:
                st.download_button("Tải tờ trình (.docx)", bb, f"to_trinh_{h.ma_ho}.docx", **ROI)
            else:
                st.caption("Cài python-docx để tải bản Word.")
            st.write("")
            st.button("✅ Duyệt", **ROI)
            st.button("❌ Từ chối", **ROI)
    with t2:
        st.markdown("**Doanh thu hóa đơn và dòng tiền ngân hàng** (triệu đồng)")
        d12 = pd.DataFrame({"Tháng": [f"T{i}" for i in range(1, 13)],
                            "Doanh thu hóa đơn": np.array(h.doanh_thu_hoa_don) / 1e6,
                            "Dòng tiền ngân hàng": np.array(h.dong_tien_ngan_hang) / 1e6})
        st.altair_chart(bd_duong(d12, "Tháng", None, 300, mau=[TEAL, AMBER]), **ROI)
        e1, e2 = st.columns([3, 2])
        with e1:
            st.markdown("**Tiền điện theo tháng** (triệu đồng)")
            dd = pd.DataFrame({"Tháng": [f"T{i}" for i in range(1, 13)],
                               "Tiền điện": (np.array(h.tien_dien) / 1e6).round(2)})
            st.altair_chart(bd_cot(dd, "Tháng", "Tiền điện", AMBER, 230, dinh_dang=".1f"), **ROI)
        with e2:
            st.markdown("**Đối chiếu chéo tiền điện**")
            chuan = DIEN_CHUAN_NGANH.get(h.nganh, 0.025)
            thuc = sum(h.tien_dien) / max(sum(h.doanh_thu_hoa_don), 1)
            st.metric("Tỷ lệ tiền điện / doanh thu", f"{thuc:.2%}",
                      f"{d['dien_lech_nganh']:+.0%} so với chuẩn ngành {chuan:.2%}")
            st.markdown(f"<div class='note'>{dien_giai('dien_lech_nganh', d['dien_lech_nganh'])}"
                        "<br><br>Đây là cơ chế chống thổi khống hóa đơn: điện khó làm giả hơn "
                        "doanh thu.</div>", unsafe_allow_html=True)
    with t3:
        v1, v2 = st.columns([3, 2])
        with v1:
            st.markdown("**Đóng góp của từng yếu tố vào điểm rủi ro**")
            s = pd.DataFrame([(TEN_DAC_TRUNG.get(k, k), round(v, 4)) for k, v in tt.dong_gop_shap],
                             columns=["Yếu tố", "Đóng góp"])
            st.altair_chart(bd_cot_am_duong(s, "Yếu tố", "Đóng góp", 360), **ROI)
            st.caption("Cam làm TĂNG rủi ro, xanh làm GIẢM. Với mô hình tuyến tính, đóng góp "
                       "tính chính xác bằng hệ số × giá trị chuẩn hóa.")
        with v2:
            st.markdown("**Bộ 11 đặc trưng của hồ sơ này**")
            st.dataframe(pd.DataFrame([{"Đặc trưng": TEN_DAC_TRUNG[k], "Giá trị": round(d[k], 4)}
                                       for k in COT]), hide_index=True, height=360, **ROI)
    with t4:
        k = tt.ket_qua_kiem_soat
        m1, m2, m3 = st.columns(3)
        kpi(m1, "Tỷ lệ truy vết", f"{k['ty_le_truy_vet']:.0%}",
            "mọi con số về được bằng chứng" if k["dat"] else "có con số cần kiểm tra tay", hi=not k["dat"])
        kpi(m2, "Số bằng chứng", f"{k['so_bang_chung']}", "mẩu dữ liệu có nguồn")
        kpi(m3, "Con số trong tờ trình", f"{k['tong_so_con_so']}", "đã đối chiếu từng số")
        st.write("")
        if k["so_khong_truy_duoc"]:
            st.warning("Chưa truy được nguồn: " + ", ".join(k["so_khong_truy_duoc"][:12]))
        with st.expander("Danh sách bằng chứng", expanded=True):
            st.dataframe(pd.DataFrame([{"Mã": x.ma, "Nội dung": x.noi_dung, "Nguồn": x.nguon}
                                       for x in tt.bang_chung]), hide_index=True, **ROI)
        with st.expander("Nhật ký tác tử"):
            for dg in tt.nhat_ky: st.text(dg)
        with st.expander("Dữ kiện đưa cho tác tử 4"):
            st.text("\n".join(tt.du_kien))

# ══════════════════════════════════════════════ 3. NHẬP HỒ SƠ MỚI
elif man == "Nhập hồ sơ mới":
    st.title("Nhập hồ sơ mới")
    st.markdown("<div class='note'>Nhập số liệu của một hộ kinh doanh bất kỳ để hệ thống chấm điểm. "
                "Chọn một hồ sơ mẫu ở dưới để điền sẵn rồi sửa lại cho nhanh.</div>",
                unsafe_allow_html=True)
    st.write("")
    mau = st.selectbox("Điền sẵn từ hồ sơ mẫu", df.ma_ho.tolist(),
                       format_func=lambda m: f"{m} — {df[df.ma_ho==m].ten_ho.iloc[0]}")
    g = df[df.ma_ho == mau].iloc[0]
    with st.form("ho_so_moi"):
        c1, c2, c3 = st.columns(3)
        ten = c1.text_input("Tên hộ kinh doanh", value=str(g.ten_ho) + " (bản sửa)")
        ds_nganh = sorted(set(DIEN_CHUAN_NGANH) | set(df["nganh"].dropna().unique()))
        nganh = c2.selectbox("Ngành", ds_nganh,
                             index=ds_nganh.index(g.nganh) if g.nganh in ds_nganh else 0)
        tinh = c3.text_input("Tỉnh/Thành phố", value=str(g.tinh_thanh))
        c4, c5, c6, c7 = st.columns(4)
        nam = c4.number_input("Số năm hoạt động", 0.1, 40.0, float(g.so_nam_hoat_dong), 0.1)
        gd = c5.number_input("Giao dịch trung bình/ngày", 1, 2000, int(g.so_giao_dich_trung_binh_ngay))
        cic = c6.selectbox("Có lịch sử CIC", ["Có", "Không"],
                           index=0 if str(g.co_lich_su_tin_dung_CIC).strip() == "Có" else 1)
        vay = c7.number_input("Số tiền đề nghị vay (triệu đồng)", 5.0, 5000.0,
                              float(g.so_tien_de_nghi_vay) / 1e6, 5.0)
        st.markdown("**Số liệu 12 tháng** (đơn vị: triệu đồng)")
        b12 = pd.DataFrame({
            "Tháng": [f"T{i}" for i in range(1, 13)],
            "Doanh thu hóa đơn": [float(g[f"doanh_thu_hoa_don_T{i}"]) / 1e6 for i in range(1, 13)],
            "Dòng tiền ngân hàng": [float(g[f"dong_tien_ngan_hang_T{i}"]) / 1e6 for i in range(1, 13)],
            "Tiền điện": [float(g[f"tien_dien_T{i}"]) / 1e6 for i in range(1, 13)]})
        sua = st.data_editor(b12, hide_index=True, disabled=["Tháng"], height=460, **ROI,
                             column_config={c: st.column_config.NumberColumn(format="%.3f")
                                            for c in b12.columns[1:]})
        gui = st.form_submit_button("Chấm điểm hồ sơ này", type="primary")

    if gui:
        hs = HoSo(ma_ho="HKD-MOI", ten_ho=ten, nganh=nganh, tinh_thanh=tinh,
                  so_nam_hoat_dong=float(nam),
                  doanh_thu_hoa_don=[float(x) * 1e6 for x in sua["Doanh thu hóa đơn"]],
                  dong_tien_ngan_hang=[float(x) * 1e6 for x in sua["Dòng tiền ngân hàng"]],
                  tien_dien=[float(x) * 1e6 for x in sua["Tiền điện"]],
                  so_giao_dich_ngay=int(gd), co_cic=(cic == "Có"),
                  so_tien_de_nghi_vay=float(vay) * 1e6)
        thanh = st.progress(0.0, "Bắt đầu…")
        tt = tham_dinh(hs, dung_llm=dung_llm_ct,
                       bao_tien_trinh=lambda i, n, t: thanh.progress(i / n, f"Tác tử {t}"))
        thanh.empty()
        st.success("Đã chạy xong sáu tác tử.")
        c = st.columns(4)
        kpi(c[0], "Xác suất vỡ nợ", f"{tt.xac_suat_vo_no:.1%}", "trong 12 tháng", hi=True)
        kpi(c[1], "Hạng rủi ro", tt.hang_rui_ro, "theo ngưỡng chính sách")
        kpi(c[2], "Đề xuất duyệt", tien(tt.khuyen_nghi["de_xuat_duyet"]), "triệu đồng")
        kpi(c[3], "Tỷ lệ truy vết", f"{tt.ket_qua_kiem_soat['ty_le_truy_vet']:.0%}", "số có nguồn")
        st.write("")
        x1, x2 = st.columns([3, 2])
        x1.text_area("tt", tt.to_trinh, height=430, label_visibility="collapsed")
        with x2:
            st.markdown("**Yếu tố ảnh hưởng**")
            s = pd.DataFrame([(TEN_DAC_TRUNG.get(k, k), round(v, 4))
                              for k, v in tt.dong_gop_shap[:8]], columns=["Yếu tố", "Đóng góp"])
            st.altair_chart(bd_cot_am_duong(s, "Yếu tố", "Đóng góp", 340), **ROI)
        st.download_button("Tải tờ trình (.txt)", tt.to_trinh.encode("utf-8"), "to_trinh_ho_so_moi.txt")

# ══════════════════════════════════════════════ 4. NẠP DỮ LIỆU MỚI
elif man == "Nạp dữ liệu mới":
    st.title("Nạp dữ liệu mới")
    st.markdown("<div class='note'>Tải lên bộ hồ sơ của bạn để hệ thống chấm điểm toàn bộ. "
                "Nếu tệp có cột <b>nhan_vo_no</b> (0 hoặc 1) thì huấn luyện lại được mô hình "
                "và so sánh trực tiếp với mô hình hiện tại.</div>", unsafe_allow_html=True)
    st.write("")

    c1, c2 = st.columns([3, 2])
    with c1:
        st.markdown(f"**Bộ dữ liệu đang dùng:** {ten_nguon()} · {len(df)} hồ sơ")
    with c2:
        st.download_button("Tải tệp mẫu (.csv)",
                           nap_mac_dinh().head(3).to_csv(index=False).encode("utf-8-sig"),
                           "micra_tep_mau.csv", "text/csv", **ROI,
                           help="Ba dòng mẫu với đầy đủ cột đúng tên. Xóa dữ liệu, giữ hàng "
                                "tiêu đề, rồi điền hồ sơ của bạn vào.")

    with st.expander("Yêu cầu về tệp"):
        st.markdown(
            f"- Định dạng **.csv** (mã hóa UTF-8) hoặc **.xlsx**\n"
            f"- Tối thiểu **5 hồ sơ**, mỗi hồ sơ một dòng\n"
            f"- Bắt buộc có **{len(COT_BAT_BUOC)} cột**, trong đó 36 cột số liệu 12 tháng\n"
            f"- Đơn vị tiền: **đồng**, không phải triệu đồng\n"
            f"- Cột `co_lich_su_tin_dung_CIC` nhận giá trị `Có` hoặc `Không`\n"
            f"- Cột `nhan_vo_no` không bắt buộc; có thì mới huấn luyện lại được")
        st.code(", ".join(COT_BAT_BUOC[:8]) + ",\n"
                + "doanh_thu_hoa_don_T1 … doanh_thu_hoa_don_T12,\n"
                + "dong_tien_ngan_hang_T1 … dong_tien_ngan_hang_T12,\n"
                + "tien_dien_T1 … tien_dien_T12,\n"
                + "nhan_vo_no  (không bắt buộc)", language=None)

    tep = st.file_uploader("Chọn tệp dữ liệu", type=["csv", "xlsx", "xls"])

    if tep is not None:
        try:
            moi = (pd.read_csv(tep) if tep.name.lower().endswith(".csv")
                   else pd.read_excel(tep))
        except Exception as e:
            st.error(f"Không đọc được tệp: {e}")
            moi = None

        if moi is not None:
            loi, canh_bao = kiem_tra_bang(moi)
            k1, k2, k3 = st.columns(3)
            kpi(k1, "Số hồ sơ", f"{len(moi)}", f"{len(moi.columns)} cột")
            kpi(k2, "Kết quả kiểm tra", "ĐẠT" if not loi else "CHƯA ĐẠT",
                f"{len(loi)} lỗi · {len(canh_bao)} cảnh báo", hi=bool(loi))
            co_nhan = "nhan_vo_no" in moi.columns and not loi
            kpi(k3, "Huấn luyện lại", "Được" if co_nhan else "Không",
                f"{int(moi['nhan_vo_no'].sum())} ca vỡ nợ" if co_nhan else "thiếu cột nhan_vo_no")
            st.write("")

            for x in loi:
                st.error(x)
            for x in canh_bao:
                st.warning(x)

            st.markdown("**Xem trước 10 dòng đầu**")
            st.dataframe(moi.head(10), hide_index=True, height=240, **ROI)

            if not loi:
                st.divider()
                a, b_ = st.columns(2)
                with a:
                    st.markdown("#### Bước 1 — Dùng bộ dữ liệu này")
                    st.caption("Toàn bộ màn hình khác sẽ chuyển sang chạy trên bộ dữ liệu vừa tải. "
                               "Mô hình giữ nguyên.")
                    if st.button("Nạp và chấm điểm", type="primary", **ROI):
                        st.session_state["bang_moi"] = moi
                        st.session_state["ten_nguon"] = tep.name
                        st.cache_resource.clear()
                        st.success(f"Đã nạp {len(moi)} hồ sơ từ {tep.name}.")
                        st.rerun()
                with b_:
                    st.markdown("#### Bước 2 — Huấn luyện lại mô hình")
                    if not co_nhan:
                        st.caption("Tệp không có cột nhan_vo_no nên bước này bị khóa.")
                        st.button("Huấn luyện lại", disabled=True, **ROI)
                    elif int(moi["nhan_vo_no"].sum()) < 5 or len(moi) < 30:
                        st.caption(f"Cần tối thiểu 30 hồ sơ và 5 ca vỡ nợ. "
                                   f"Tệp đang có {len(moi)} hồ sơ và "
                                   f"{int(moi['nhan_vo_no'].sum())} ca vỡ nợ.")
                        st.button("Huấn luyện lại", disabled=True, **ROI)
                    else:
                        st.caption("Mô hình cũ được sao lưu tự động trước khi ghi đè.")
                        if st.button("Huấn luyện lại trên dữ liệu này", **ROI):
                            from train import huan_luyen
                            goc = MODELS / "mo_hinh_rui_ro.joblib"
                            luu = MODELS / f"mo_hinh_sao_luu_{datetime.now():%Y%m%d_%H%M%S}.joblib"
                            if goc.exists():
                                shutil.copy2(goc, luu)
                                st.session_state["sao_luu"] = str(luu)
                            with st.spinner("Đang huấn luyện… mất khoảng 10 đến 30 giây"):
                                _, kq = huan_luyen(luu=True, im_lang=True, bang=moi)
                            scoring.nap_lai()
                            st.cache_resource.clear(); st.cache_data.clear()
                            st.session_state["ket_qua_hl"] = {
                                "chon": kq["_chon"], "n": kq["_so_mau"], "npos": kq["_so_vo_no"],
                                "auc": kq[kq["_chon"]]["auc"], "pr": kq[kq["_chon"]]["pr_auc"],
                                "brier": kq[kq["_chon"]]["brier"]}
                            st.rerun()

    if "ket_qua_hl" in st.session_state:
        r = st.session_state["ket_qua_hl"]
        st.divider()
        st.markdown("### Kết quả huấn luyện lại")
        c = st.columns(4)
        kpi(c[0], "Mô hình được chọn", r["chon"].upper(), "chọn bằng kiểm định chéo")
        kpi(c[1], "ROC-AUC mới", f"{r['auc'][0]:.3f}", f"± {r['auc'][1]:.3f}", hi=True)
        kpi(c[2], "PR-AUC", f"{r['pr'][0]:.3f}", f"± {r['pr'][1]:.3f}")
        kpi(c[3], "Cỡ mẫu", f"{r['n']}", f"{r['npos']} ca vỡ nợ")
        st.markdown(f"<div class='warn'>So với mô hình gốc (AUC 0,858 trên 60 hồ sơ mô phỏng): "
                    f"chênh lệch <b>{r['auc'][0]-0.858:+.3f}</b>. Nếu độ chênh nhỏ hơn độ lệch "
                    f"chuẩn thì chưa kết luận được mô hình nào tốt hơn.</div>", unsafe_allow_html=True)
        if "sao_luu" in st.session_state:
            st.caption(f"Mô hình cũ đã lưu tại: {Path(st.session_state['sao_luu']).name}")

    st.divider()
    if st.button("Khôi phục bộ dữ liệu và mô hình gốc"):
        st.session_state.pop("bang_moi", None)
        st.session_state.pop("ten_nguon", None)
        st.session_state.pop("ket_qua_hl", None)
        sl = st.session_state.pop("sao_luu", None)
        if sl and Path(sl).exists():
            shutil.copy2(sl, MODELS / "mo_hinh_rui_ro.joblib")
        scoring.nap_lai()
        st.cache_resource.clear(); st.cache_data.clear()
        st.success("Đã quay lại bộ dữ liệu mô phỏng 60 hộ và mô hình gốc.")
        st.rerun()

# ══════════════════════════════════════════════ 5. SO SÁNH
elif man == "So sánh hồ sơ":
    st.title("So sánh hồ sơ")
    st.caption("Đặt cạnh nhau để thấy vì sao hai hộ có doanh thu tương đương lại được xếp hạng khác nhau.")
    tat_ca = chay_tat_ca()
    chon = st.multiselect("Chọn từ hai đến bốn hồ sơ", df.ma_ho.tolist(),
                          default=df.ma_ho.tolist()[:2], max_selections=4,
                          format_func=lambda m: f"{m} — {df[df.ma_ho==m].ten_ho.iloc[0]}")
    if len(chon) < 2:
        st.info("Chọn thêm ít nhất một hồ sơ nữa.")
    else:
        for c, m in zip(st.columns(len(chon)), chon):
            t = tat_ca[m]
            with c:
                st.markdown(f"**{t.ho_so.ten_ho}**")
                st.markdown(chip(t.ho_so.nganh, TEAL) + chip(t.hang_rui_ro, MAU[t.hang_rui_ro]),
                            unsafe_allow_html=True)
                st.metric("Xác suất vỡ nợ", f"{t.xac_suat_vo_no:.1%}")
                st.metric("Doanh thu TB", tien(t.dac_trung["dt_trung_binh"]) + " đ")
                st.metric("Đề xuất duyệt", tien(t.khuyen_nghi["de_xuat_duyet"]) + " đ")
        st.divider()
        st.markdown("**Xác suất vỡ nợ**")
        dp = pd.DataFrame({"Hộ": [tat_ca[m].ho_so.ten_ho for m in chon],
                           "Xác suất %": [round(tat_ca[m].xac_suat_vo_no * 100, 1) for m in chon]})
        st.altair_chart(bd_cot(dp, "Hộ", "Xác suất %", TEAL, 220, ngang=True, dinh_dang=".1f"), **ROI)
        st.markdown("**Bảng đặc trưng đối chiếu**")
        ss = pd.DataFrame({tat_ca[m].ho_so.ten_ho: {TEN_DAC_TRUNG[k]: round(tat_ca[m].dac_trung[k], 4)
                                                    for k in COT} for m in chon})
        st.dataframe(ss, height=420, **ROI)
        st.markdown("**Doanh thu hóa đơn theo tháng** (triệu đồng)")
        dl = pd.DataFrame({"Tháng": [f"T{i}" for i in range(1, 13)]})
        for m in chon:
            dl[tat_ca[m].ho_so.ten_ho] = np.array(tat_ca[m].ho_so.doanh_thu_hoa_don) / 1e6
        st.altair_chart(bd_duong(dl, "Tháng", None, 300), **ROI)

# ══════════════════════════════════════════════ 6. GIÁM SÁT
elif man == "Giám sát sau vay":
    st.title("Giám sát sau giải ngân")
    st.caption("Tác tử 6 chạy hằng tuần trên dữ liệu mới nhất. Mục tiêu: phát hiện sớm trước khi thành nợ xấu.")
    mo_phong = st.toggle("Mô phỏng thêm 3 tháng dữ liệu sau giải ngân", value=True)
    if mo_phong:
        st.markdown("<div class='warn'>Ba tháng dữ liệu dưới đây là <b>mô phỏng</b>, "
                    "không phải dữ liệu thật.</div>", unsafe_allow_html=True)
        st.write("")
    tat_ca = chay_tat_ca()
    kq = []
    for ma, tt in tat_ca.items():
        g = (a6_giam_sat.chay(tt, ba_thang_tiep(tt.ho_so, tt.xac_suat_vo_no)) if mo_phong
             else tt.du_lieu_tho["giam_sat"])
        kq.append({"Mã": ma, "Hộ kinh doanh": tt.ho_so.ten_ho, "Ngành": tt.ho_so.nganh,
                   "Mức cảnh báo": g["muc_canh_bao"], "Số dấu hiệu": len(g["chi_tiet"]),
                   "Xác suất vỡ nợ": tt.xac_suat_vo_no, "Hành động đề xuất": g["hanh_dong"]})
    t = pd.DataFrame(kq)
    c = st.columns(4)
    kpi(c[0], "Cảnh báo đỏ", int((t["Mức cảnh báo"] == "Đỏ").sum()), "cần xử lý ngay", hi=True)
    kpi(c[1], "Cảnh báo vàng", int((t["Mức cảnh báo"] == "Vàng").sum()), "theo dõi sát")
    kpi(c[2], "Bình thường", int((t["Mức cảnh báo"] == "Xanh").sum()), "không dấu hiệu bất thường")
    kpi(c[3], "Chi phí giám sát", "≈ 0", "chi phí biên mỗi lần quét")
    st.write("")
    dm = (t["Mức cảnh báo"].value_counts().reindex(["Đỏ", "Vàng", "Xanh"]).fillna(0)
          .astype(int).rename_axis("Mức").reset_index(name="Số hộ"))
    cm = alt.Chart(dm).mark_bar(cornerRadius=3).encode(
        x=alt.X("Mức:N", sort=["Đỏ", "Vàng", "Xanh"], title=None),
        y=alt.Y("Số hộ:Q", title=None),
        color=alt.Color("Mức:N", legend=None, scale=alt.Scale(
            domain=["Đỏ", "Vàng", "Xanh"], range=["#C00000", "#B8860B", "#1F7A3D"])),
        tooltip=["Mức", "Số hộ"])
    st.altair_chart(_khung(cm, 200), **ROI)
    thu_tu = {"Đỏ": 0, "Vàng": 1, "Xanh": 2}
    t = t.sort_values("Mức cảnh báo", key=lambda s: s.map(thu_tu))
    st.dataframe(t, hide_index=True, height=460, **ROI,
                 column_config={"Xác suất vỡ nợ": st.column_config.ProgressColumn(
                     "Xác suất vỡ nợ", format="%.0f%%", min_value=0.0, max_value=1.0)})
    st.download_button("Tải danh sách cảnh báo (.csv)", t.to_csv(index=False).encode("utf-8-sig"),
                       "micra_canh_bao_sau_vay.csv", "text/csv")

# ══════════════════════════════════════════════ 7. PHỎNG VẤN
elif man == "Phỏng vấn qua Zalo":
    st.title("Phỏng vấn hộ kinh doanh qua Zalo")
    st.caption("Tác tử 2 — thay khâu khảo sát thực địa, nơi tập trung phần lớn chi phí thẩm định.")
    if "chat" not in st.session_state:
        st.session_state.chat = [{"role": "assistant", "content": CAU_HOI[0][1]}]
    da_hoi = sum(1 for m in st.session_state.chat if m["role"] == "assistant")
    st.progress(min(da_hoi / len(CAU_HOI), 1.0), f"Đã hỏi {min(da_hoi, len(CAU_HOI))}/{len(CAU_HOI)} câu")
    for m in st.session_state.chat:
        st.chat_message(m["role"]).write(m["content"])
    if tra_loi := st.chat_input("Trả lời với vai chủ hộ kinh doanh…"):
        st.session_state.chat.append({"role": "user", "content": tra_loi})
        st.chat_message("user").write(tra_loi)
        n = sum(1 for m in st.session_state.chat if m["role"] == "assistant")
        if n < len(CAU_HOI):
            reply = CAU_HOI[n][1]
        else:
            reply = goi_llm(HE_THONG,
                            "Tóm tắt cuộc phỏng vấn sau thành gạch đầu dòng, mỗi ý một dòng "
                            "bắt đầu bằng '-'. Chỉ dùng thông tin có trong hội thoại.\n\n"
                            + "\n".join(f"{m['role']}: {m['content']}" for m in st.session_state.chat),
                            max_tokens=500)
        st.session_state.chat.append({"role": "assistant", "content": reply})
        st.chat_message("assistant").write(reply)
    if st.button("Bắt đầu lại"):
        st.session_state.chat = [{"role": "assistant", "content": CAU_HOI[0][1]}]
        st.rerun()

# ══════════════════════════════════════════════ 8. KIỂM CHỨNG
else:
    st.title("Kiểm chứng mô hình")
    st.caption(f"Tính trực tiếp trên {ten_nguon()} khi mở trang này, không phải số chép sẵn.")
    m = chi_so_mo_hinh(st.session_state["khoa"])
    if m is None:
        st.warning("Bộ dữ liệu đang dùng không có cột nhan_vo_no nên không đánh giá được. "
                   "Vào màn hình Nạp dữ liệu mới để khôi phục bộ mặc định.")
    else:
        c = st.columns(5)
        kpi(c[0], "ROC-AUC", f"{m['auc'][0]:.3f}", f"± {m['auc'][1]:.3f} qua 10 lần chia", hi=True)
        kpi(c[1], "Khoảng tin cậy 95%", f"{m['ktc'][0]:.2f} – {m['ktc'][1]:.2f}", "bootstrap 2.000 lần")
        kpi(c[2], "PR-AUC", f"{m['pr'][0]:.3f}", f"± {m['pr'][1]:.3f}")
        kpi(c[3], "Brier score", f"{m['brier'][0]:.3f}", "càng thấp càng tốt")
        kpi(c[4], "Cỡ mẫu", f"{m['n']}", f"{m['n_vo_no']} ca vỡ nợ ({m['n_vo_no']/m['n']:.0%})")
        st.write("")
        st.markdown("<div class='note'><b>Đọc con số này thế nào.</b> AUC quanh 0,86 là mức khá "
                    "cho bài toán tín dụng, nhưng khoảng tin cậy rất rộng vì mẫu nhỏ. Nhóm báo "
                    "cáo cả khoảng thay vì một con số duy nhất, và không bao giờ báo cáo AUC "
                    "trên tập huấn luyện — con số đó luôn đẹp và luôn vô nghĩa.</div>",
                    unsafe_allow_html=True)
        st.write("")
        g1, g2 = st.columns(2)
        with g1:
            st.markdown("**Đường cong ROC** (kiểm định chéo 5 phần)")
            st.altair_chart(bd_roc(m["roc"]), **ROI)
            st.caption("Đường đứt là mức đoán ngẫu nhiên. Càng cong lên trái càng tốt.")
        with g2:
            st.markdown("**Hiệu chỉnh xác suất** — dự báo so với thực tế")
            st.altair_chart(bd_hieu_chinh(m["hieu_chinh"]), **ROI)
            st.caption("Càng bám đường chéo thì '13%' càng thật sự nghĩa là 13 trên 100 hộ.")
        st.divider()
        st.markdown("**Xác suất dự báo trung bình theo kết quả thực tế**")
        ph = pd.DataFrame({"p": m["p_cv"], "Kết quả thực": np.where(m["y"] == 1, "Vỡ nợ", "Trả đủ")})
        dg = ph.groupby("Kết quả thực")["p"].mean().mul(100).round(1).reset_index(name="Xác suất %")
        st.altair_chart(bd_cot(dg, "Kết quả thực", "Xác suất %", SEA, 210, dinh_dang=".1f"), **ROI)
        with st.expander("Vì sao không dùng mô hình phức tạp hơn"):
            st.markdown(
                "Nhóm đã thử tám thuật toán: hồi quy logistic, rừng ngẫu nhiên, Extra Trees, "
                "Gradient Boosting, XGBoost, SVM, RBF, Naive Bayes. Phương án tốt nhất chỉ hơn "
                "**0,003 điểm AUC** — nhỏ hơn nhiều so với độ rộng khoảng tin cậy, tức là không "
                "có ý nghĩa thống kê.\n\n"
                "Nhóm cũng thử bổ sung năm đặc trưng mới. Bộ 16 đặc trưng cho AUC **0,774**, "
                "THẤP HƠN bộ 11 đặc trưng gốc.\n\n"
                "Kết luận: ràng buộc nằm ở **dữ liệu**, không nằm ở mô hình. Cần 300–500 hồ sơ "
                "thật để cải thiện thực chất, tương đương 45–75 ca vỡ nợ.")
